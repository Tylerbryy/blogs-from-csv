#make sure your CSV file has "Title" at the top

import subprocess
import importlib
import sys

def install_package(package_name):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])

REQUIRED_PACKAGES = ["openai", "tqdm", "colorama", "tkinter", "python-docx", "backoff"]

for package in REQUIRED_PACKAGES:
    try:
        importlib.import_module(package)
    except ImportError:
        install_package(package)

import openai
import os
import csv
from docx import Document
from tqdm import tqdm
from colorama import init, Fore, Style
import tkinter as tk
from tkinter import filedialog
import backoff
# Initialize colorama
init()

def clear_screen():
    os.system('cls' if os.name == 'nt' else 'clear')

clear_screen()

openai.api_key = "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"

def select_file_path():
    root = tk.Tk()
    root.withdraw()  # Hide the root window
    file_path = filedialog.askopenfilename(title="Select CSV file", filetypes=[("CSV files", "*.csv")])
    return file_path

def load_data_from_csv(file_path):
    with open(file_path, newline='', encoding='utf-8') as csvfile:  
        reader = csv.DictReader(csvfile)
        return list(reader)


@backoff.on_exception(backoff.expo, openai.error.Timeout, max_tries=100)
def generate_blog_content(title: str):
    messages = []
    messages.append({"role": "user", "content": f"""Create a blog post about '{title}'."""})

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=messages,
            timeout=1200
        )
    except openai.error.Timeout as e:
        print(f"Request timed out: {e}")
        raise  


    messages.clear()
    return response.choices[0].message.content

def sanitize_filename(filename, char_limit=50):
    # Remove invalid characters for filenames
    invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
    for char in invalid_chars:
        filename = filename.replace(char, '')
    
    # Truncate filename if it exceeds the character limit
    if len(filename) > char_limit:
        filename = filename[:char_limit]  # Keep the first 'char_limit' characters

    return filename


def generate_blogs(data):
    # Check if the folder exists. If not, create one.
    folder_name = "Generated_Blogs"
    if not os.path.exists(folder_name):
        os.mkdir(folder_name)

    titles = [doc['Title'] for doc in data]

    progress_bar = tqdm(titles, desc="Generating Blogs", unit="blog", bar_format="{desc}: {percentage:3.0f}%|{bar}| {n_fmt}/{total_fmt}", colour='green', ncols=80)

    for title in progress_bar:
        document = Document()  # Create a new document for each blog
        blog_content = generate_blog_content(title)
        document.add_heading(title, level=1)
        sentences = blog_content.split('. ')
        paragraph_text = '. '.join(sentences)
        document.add_paragraph(paragraph_text)
        
        sanitized_title = sanitize_filename(title)  # Sanitize the title to create a valid filename
        document.save(f"{folder_name}/{sanitized_title}.docx")  # Save the blog in the folder

        progress_bar.set_postfix({"Current Title": title})
        progress_bar.update(1)

def main():
    # Use the file finder function to get the file path
    print("Please select your CSV file...")
    file_path = select_file_path()

    # If no file is selected, exit the program
    if not file_path:
        print(Fore.RED + "No file selected. Exiting program." + Style.RESET_ALL)
        return

    data = load_data_from_csv(file_path)

    

    generate_blogs(data)

    clear_screen()
    print(Fore.LIGHTGREEN_EX + "Blog generation completed." + Style.RESET_ALL)

if __name__ == '__main__':
    main()
